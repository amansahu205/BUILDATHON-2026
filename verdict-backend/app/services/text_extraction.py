"""
Text extraction service for uploaded documents.

Supports PDF (via pdfplumber) and DOCX (via python-docx).
Returns structured text with page/line metadata for downstream
ingestion into Databricks Vector Search.
"""

import io
import logging
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ExtractedChunk:
    content: str
    page: int | None = None
    line: int | None = None


def extract_pdf(data: bytes) -> list[ExtractedChunk]:
    """Extract text from a PDF file, returning one chunk per page."""
    import pdfplumber

    chunks: list[ExtractedChunk] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    chunks.append(ExtractedChunk(content=text.strip(), page=i, line=1))
    except Exception as exc:
        logger.error("PDF extraction failed: %s", exc)
        raise ValueError(f"Failed to extract text from PDF: {exc}") from exc

    if not chunks:
        raise ValueError("No text content found. PDF may be image-only.")
    return chunks


def extract_docx(data: bytes) -> list[ExtractedChunk]:
    """Extract text from a DOCX file, returning chunks of grouped paragraphs."""
    from docx import Document

    chunks: list[ExtractedChunk] = []
    try:
        doc = Document(io.BytesIO(data))
        current_text: list[str] = []
        line_num = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_text:
                    chunks.append(ExtractedChunk(
                        content="\n".join(current_text),
                        page=None,
                        line=line_num,
                    ))
                    current_text = []
                continue
            current_text.append(text)
            line_num += 1

        if current_text:
            chunks.append(ExtractedChunk(
                content="\n".join(current_text),
                page=None,
                line=line_num - len(current_text),
            ))

        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                chunks.append(ExtractedChunk(
                    content="\n".join(rows_text),
                    page=None,
                    line=None,
                ))
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
        raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc

    if not chunks:
        raise ValueError("No text content found in DOCX.")
    return chunks


def extract_txt(data: bytes) -> list[ExtractedChunk]:
    """Extract text from a plain text file."""
    text = data.decode("utf-8", errors="replace").strip()
    if not text:
        raise ValueError("No text content found in file.")
    return [ExtractedChunk(content=text, page=1, line=1)]


def extract_text(data: bytes, mime_type: str) -> list[ExtractedChunk]:
    """Route to the appropriate extractor based on MIME type."""
    if mime_type == "application/pdf":
        return extract_pdf(data)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return extract_docx(data)
    elif mime_type.startswith("text/"):
        return extract_txt(data)
    else:
        raise ValueError(f"Unsupported file type: {mime_type}")
